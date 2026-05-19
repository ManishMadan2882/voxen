import io
from typing import Iterable

from docx import Document as DocxDocument

from rag.document_model import DocType
from rag.ingestors.base import FileIngestor, ParsedSegment


class DocxIngestor(FileIngestor):
    extensions = (".docx",)
    doc_type = DocType.docx

    def parse(self, content: bytes, filename: str) -> Iterable[ParsedSegment]:
        doc = DocxDocument(io.BytesIO(content))
        parts: list[str] = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        text = "\n".join(parts).strip()
        if text:
            yield ParsedSegment(source=filename, page=1, text=text)
